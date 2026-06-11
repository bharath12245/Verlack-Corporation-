import os
import sys
import subprocess

def install_and_import(package, import_name=None):
    if import_name is None:
        import_name = package
    try:
        __import__(import_name)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# Ensure required libraries are installed
install_and_import("python-docx", "docx")
install_and_import("pypdf")

import docx
import pypdf

os.makedirs("extracted_content", exist_ok=True)

def extract_docx(file_path, output_path):
    print(f"Extracting DOCX: {file_path} -> {output_path}")
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Table extraction
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(full_text))

def extract_pdf(file_path, output_path):
    print(f"Extracting PDF: {file_path} -> {output_path}")
    reader = pypdf.PdfReader(file_path)
    full_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        full_text.append(f"--- Page {i+1} ---")
        full_text.append(text)
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(full_text))

if __name__ == "__main__":
    # Check workspace files
    files = os.listdir(".")
    for f in files:
        if f.endswith(".docx"):
            extract_docx(f, os.path.join("extracted_content", f.replace(".docx", ".txt")))
        elif f.endswith(".pdf"):
            extract_pdf(f, os.path.join("extracted_content", f.replace(".pdf", ".txt")))
    print("Extraction completed!")
